from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Semestral" / "documentacion" / "Guia_Explicacion_Funcional_Planilla_Prospera.docx"
LOGO = ROOT / "Semestral" / "assets" / "img" / "logo.png"

AZUL = "16223F"
DORADO = "C9A227"
GRIS = "F1F3F8"
VERDE = "E5F4EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_bullet(document, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(text, style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(text, style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(26)
styles["Title"].font.color.rgb = RGBColor.from_string(AZUL)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor.from_string(AZUL)

# Portada
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(1.25))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
r = p.add_run("GUÍA DE EXPLICACIÓN FUNCIONAL")
r.bold = True
r.font.name = "Aptos Display"
r.font.size = Pt(24)
r.font.color.rgb = RGBColor.from_string(AZUL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Planillas — La Prospera, S.A.")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor.from_string(DORADO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Examen semestral de Contabilidad | Grupo 5").italic = True

doc.add_paragraph("")
table = doc.add_table(rows=5, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, text in enumerate([
    "Alberto Chen — Base de datos y registro de personal",
    "Manuel Borbua — Cálculo central de planilla",
    "Brian Cona — Validación de los casos del Grupo 5",
    "Arturo Rodríguez — Reportes e impresión",
    "Daniela Anaya — Interfaz, navegación y correo",
]):
    table.cell(i, 0).text = text
    table.cell(i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(table.cell(i, 0), GRIS if i % 2 == 0 else "FFFFFF")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.add_run("Periodo trabajado: segunda quincena de junio de 2026").bold = True

doc.add_page_break()

add_heading(doc, "1. ¿Qué hace el programa?", 1)
doc.add_paragraph(
    "Planilla Prospera es una aplicación web que automatiza el proceso de elaboración de una "
    "planilla quincenal. El sistema registra los datos de los colaboradores, recibe las novedades "
    "del periodo, calcula ingresos y deducciones, guarda los resultados y genera los reportes "
    "solicitados por la profesora."
)

add_bullet(doc, "Registra y consulta la información del personal.")
add_bullet(doc, "Busca al colaborador antes de calcular su planilla.")
add_bullet(doc, "Calcula salario quincenal, bonificación, horas extra, dietas y comisiones.")
add_bullet(doc, "Calcula Seguro Social, Seguro Educativo, ISR y descuentos personales.")
add_bullet(doc, "Guarda cada resultado para mantener un historial.")
add_bullet(doc, "Genera reporte individual, reporte grupal y reporte para la CSS.")
add_bullet(doc, "Permite imprimir los reportes y prepara su envío por correo.")

add_heading(doc, "2. Flujo funcional del sistema", 1)
for step in [
    "El usuario entra a la pantalla de inicio y observa un resumen del sistema.",
    "En Colaboradores puede consultar, buscar, registrar o editar al personal.",
    "En Calcular planilla selecciona a un colaborador o procesa automáticamente los cuatro casos del Grupo 5.",
    "El sistema obtiene el salario base guardado y aplica las novedades del periodo.",
    "La lógica central calcula los ingresos, deducciones y salario neto.",
    "El resultado se guarda en la base de datos sin duplicar la planilla del mismo periodo.",
    "Desde el historial se consulta cualquier resultado guardado.",
    "Los reportes presentan la información individual, grupal y para la Caja de Seguro Social.",
    "Cada reporte cuenta con una opción de impresión y una opción de envío por correo.",
]:
    add_number(doc, step)

add_heading(doc, "3. Explicación de cada pantalla", 1)

add_heading(doc, "3.1 Inicio o dashboard", 2)
doc.add_paragraph(
    "Es la pantalla principal. Presenta el número de colaboradores registrados, la cantidad de "
    "planillas calculadas para el periodo y accesos rápidos a las funciones más importantes."
)
doc.add_paragraph(
    "Cómo explicarlo: “Esta pantalla funciona como un panel de control. Desde aquí podemos conocer "
    "el estado general del sistema y entrar directamente a colaboradores, cálculo y reportes.”"
)

add_heading(doc, "3.2 Módulo de colaboradores", 2)
doc.add_paragraph(
    "Administra los datos del personal. Permite listar, buscar por nombre o cédula, registrar nuevos "
    "colaboradores, editar sus datos y eliminarlos cuando no tienen planillas relacionadas."
)
add_bullet(doc, "Nombre completo y cédula.")
add_bullet(doc, "Estado civil y tipo de declaración.")
add_bullet(doc, "Cargo y salario base mensual.")
add_bullet(doc, "Año de inicio de labores.")

add_heading(doc, "3.3 Cálculo de planilla", 2)
doc.add_paragraph(
    "Esta pantalla recibe las novedades de la quincena: dieta mensual, horas extra, ventas, "
    "porcentaje de comisión, otros ingresos y descuentos personales. La bonificación general de "
    "B/.120 se aplica automáticamente."
)
doc.add_paragraph(
    "También incluye el botón “Procesar los 4 casos del Grupo 5”, que utiliza los datos específicos "
    "del documento para calcular de una vez las planillas de Manuel, José, Federico y Estefanía."
)

add_heading(doc, "3.4 Resultado e historial", 2)
doc.add_paragraph(
    "Después del cálculo se presenta un resumen de ingresos, deducciones y salario neto. El historial "
    "conserva las planillas guardadas y permite abrir el reporte individual de cada colaborador."
)

add_heading(doc, "3.5 Reporte individual", 2)
doc.add_paragraph(
    "Muestra el comprobante de un solo colaborador: datos personales, salario quincenal, bonificación, "
    "horas extra, comisión, dieta, salario bruto, deducciones y salario neto."
)

add_heading(doc, "3.6 Reporte grupal", 2)
doc.add_paragraph(
    "Consolida en una tabla a los cuatro colaboradores del Grupo 5. Permite comparar salario base, "
    "otros ingresos, salario bruto, total de descuentos y salario neto, e incluye totales generales."
)

add_heading(doc, "3.7 Reporte para la CSS", 2)
doc.add_paragraph(
    "Presenta el nombre, cédula, salario sujeto a cotización, Seguro Social del trabajador, Seguro "
    "Educativo y total reportado. El sistema también conserva los aportes patronales de Seguro Social, "
    "Seguro Educativo y riesgo profesional."
)

add_heading(doc, "3.8 Impresión y correo", 2)
doc.add_paragraph(
    "El botón Imprimir abre la función de impresión del navegador con una vista limpia, sin menú ni "
    "botones. El botón Enviar por correo prepara el reporte en formato HTML y solicita el correo del "
    "destinatario. El envío real requiere configurar una cuenta SMTP en el servidor."
)

doc.add_page_break()

add_heading(doc, "4. ¿Cómo calcula la planilla?", 1)
calc_table = doc.add_table(rows=1, cols=3)
calc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
calc_table.style = "Table Grid"
headers = ["Concepto", "Fórmula o tratamiento", "Dato aplicado"]
for i, header in enumerate(headers):
    calc_table.cell(0, i).text = header
    set_cell_shading(calc_table.cell(0, i), AZUL)
    set_cell_text_color(calc_table.cell(0, i), "FFFFFF")

rows = [
    ("Salario quincenal", "Salario mensual ÷ 2", "Pago correspondiente a la quincena"),
    ("Bonificación", "Monto fijo del periodo", "B/.120.00 para todos"),
    ("Valor por hora", "Salario mensual ÷ 195", "Jornada de 45 horas semanales"),
    ("Horas extra diurnas", "Valor hora × cantidad × 1.25", "Recargo diurno del 25%"),
    ("Comisión", "Ventas × porcentaje", "Estefanía: 55,000 × 1.5%"),
    ("Dieta", "Dieta mensual ÷ 2", "Manuel recibe B/.300"),
    ("Salario bruto", "Suma de todos los ingresos", "Base general antes de descuentos"),
    ("Seguro Social", "Base cotizable × 9.75%", "Deducción del trabajador"),
    ("Seguro Educativo", "Base cotizable × 1.25%", "Deducción del trabajador"),
    ("ISR", "Solo si la renta anual supera B/.11,000", "Variables de esta quincena se suman una sola vez"),
    ("Otros descuentos", "Descuento mensual ÷ 2", "Ahorro o mueblería"),
    ("Salario neto", "Salario bruto − total de descuentos", "Monto final que recibe el colaborador"),
]
for row in rows:
    cells = calc_table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

add_heading(doc, "5. Casos y resultados del Grupo 5", 1)
results = doc.add_table(rows=1, cols=7)
results.style = "Table Grid"
results.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, text in enumerate(["Colaborador", "Bruto", "CSS", "S. Educ.", "ISR", "Otros desc.", "Neto"]):
    results.cell(0, i).text = text
    set_cell_shading(results.cell(0, i), AZUL)
    set_cell_text_color(results.cell(0, i), "FFFFFF")

data = [
    ("Manuel Peña", "765.00", "66.18", "8.48", "28.06", "125.00", "537.28"),
    ("José Martínez", "468.60", "45.69", "5.86", "0.00", "25.00", "392.05"),
    ("Federico Montiel", "616.15", "60.07", "7.70", "0.41", "100.00", "447.97"),
    ("Estefanía Sousa Rincón", "1,272.60", "124.08", "15.91", "0.00", "25.00", "1,107.61"),
]
for row in data:
    cells = results.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
    set_cell_shading(cells[-1], VERDE)

add_heading(doc, "6. Guion sugerido para la exposición", 1)
script = (
    "“Nuestro proyecto se llama Planilla Prospera y fue desarrollado para automatizar la planilla "
    "de la empresa La Prospera, S.A. correspondiente a la segunda quincena de junio de 2026. "
    "Funcionalmente, el sistema comienza con el registro de colaboradores. Cada colaborador tiene "
    "guardados sus datos personales, cargo, salario mensual y tipo de declaración. Cuando se calcula "
    "la planilla, el sistema busca esos datos y permite ingresar las novedades de la quincena, como "
    "horas extra, dieta, ventas, comisión y descuentos personales. Luego calcula automáticamente el "
    "salario quincenal, la bonificación de 120 balboas, el salario bruto, el Seguro Social, el Seguro "
    "Educativo, el impuesto sobre la renta y el salario neto. El resultado queda guardado en un "
    "historial y puede consultarse mediante tres reportes: individual, grupal y para la Caja de Seguro "
    "Social. Finalmente, los reportes se pueden imprimir y el sistema deja preparado su envío por "
    "correo. Para demostrarlo, procesaremos los cuatro casos del Grupo 5 y mostraremos que cada "
    "resultado se conserva en la base de datos.”"
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.paragraph_format.right_indent = Inches(0.25)
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
r = p.add_run(script)
r.italic = True

add_heading(doc, "7. Orden recomendado para la demostración", 1)
for step in [
    "Abrir Inicio y explicar los indicadores del panel.",
    "Entrar a Colaboradores y mostrar los cuatro trabajadores del Grupo 5.",
    "Usar la búsqueda con un nombre o una cédula.",
    "Entrar a Calcular planilla y pulsar “Procesar los 4 casos del Grupo 5”.",
    "Explicar un resultado, distinguiendo ingresos, deducciones y salario neto.",
    "Abrir el historial para demostrar que la información quedó guardada.",
    "Mostrar el reporte individual de un colaborador.",
    "Mostrar el reporte grupal y sus totales.",
    "Mostrar el reporte CSS y la base sujeta a cotización.",
    "Pulsar Imprimir para enseñar la vista preparada para papel o PDF.",
    "Explicar que el envío por correo queda disponible al configurar el servidor SMTP.",
]:
    add_number(doc, step)

add_heading(doc, "8. Preguntas probables y respuestas", 1)
qa = [
    ("¿Por qué el salario mínimo es B/.655.20?",
     "Porque el decreto suministrado fija B/.3.36 por hora para fabricación de cemento o concreto en la Región 1. Al multiplicarlo por 45 horas semanales, 52 semanas y dividirlo entre 12 meses se obtiene B/.655.20."),
    ("¿Cómo evitan duplicar una planilla?",
     "La base de datos tiene una restricción que permite una sola planilla por colaborador y periodo. Si se vuelve a procesar, se actualiza el resultado existente."),
    ("¿Por qué la dieta de Manuel no cotiza completa?",
     "La teoría indica que una dieta recurrente integra salario para la CSS únicamente en la parte que excede el 25% de un salario mensual."),
    ("¿Todos los colaboradores pagan impuesto sobre la renta?",
     "No. La tarifa es 0% mientras la renta neta gravable anual proyectada no supere B/.11,000. En estos casos José y Estefanía no pagan ISR; Manuel sí paga y Federico supera el límite por un monto pequeño."),
    ("¿Dónde se realizan los cálculos?",
     "En funciones centrales reutilizables. Los reportes solo consultan resultados guardados; no vuelven a calcularlos."),
    ("¿Qué información queda guardada?",
     "Los componentes de ingresos, deducciones, base cotizable, salario bruto, salario neto y aportes patronales."),
    ("¿El correo ya funciona?",
     "La pantalla y el contenido del envío están implementados. Para enviar fuera del equipo debe configurarse una cuenta SMTP en el servidor."),
    ("¿El sistema incluye otros grupos?",
     "No. La base, los casos de prueba y los reportes corresponden exclusivamente al Grupo 5."),
]
for question, answer in qa:
    p = doc.add_paragraph()
    p.add_run(question).bold = True
    p = doc.add_paragraph(answer)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)

add_heading(doc, "9. Conclusión", 1)
doc.add_paragraph(
    "El sistema cumple funcionalmente con el ciclo solicitado: captura datos del personal, busca al "
    "colaborador, calcula la planilla, guarda los resultados, presenta los reportes, permite imprimirlos "
    "y prepara su envío por correo. Además, mantiene una separación clara entre los datos, los cálculos "
    "y la presentación, lo que facilita verificar cada resultado del Grupo 5."
)

# Encabezado y pie
for section in doc.sections:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Planilla Prospera | Grupo 5")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(91, 100, 120)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Guía de explicación funcional — Examen semestral de Contabilidad")

doc.core_properties.title = "Guía de explicación funcional — Planilla Prospera"
doc.core_properties.subject = "Sistema de planillas del Grupo 5"
doc.core_properties.author = "Grupo 5"
doc.save(OUTPUT)
print(OUTPUT)
